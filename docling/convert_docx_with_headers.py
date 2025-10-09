from docx import Document
from docling.document_converter import DocumentConverter

# Path to your DOCX file
docx_file = r"C:\Users\Admin\Desktop\demo\docling\docling\cli\test.docx"

# --- Step 1: Extract headers and footers ---
doc = Document(docx_file)

all_headers = []
all_footers = []

for section in doc.sections:
    # Extract all types of headers
    for header_type in [section.header, section.first_page_header, section.even_page_header]:
        if header_type is None:
            continue
        for para in header_type.paragraphs:
            text = para.text.strip()
            if text:
                all_headers.append(text)
    
    # Extract all types of footers
    for footer_type in [section.footer, section.first_page_footer, section.even_page_footer]:
        if footer_type is None:
            continue
        for para in footer_type.paragraphs:
            text = para.text.strip()
            if text:
                all_footers.append(text)

headers_text = "\n".join(all_headers)
footers_text = "\n".join(all_footers)

# --- Step 2: Convert main body using Docling ---
converter = DocumentConverter()
docling_doc = converter.convert(docx_file).document
body_md = docling_doc.export_to_markdown()

# --- Step 3: Combine headers, body, and footers ---
final_md = ""
if headers_text:
    final_md += headers_text + "\n\n"
final_md += body_md
if footers_text:
    final_md += "\n\n" + footers_text

# --- Step 4: Save to output.md ---
output_path = "output.md"  # This will create output.md in the current folder
try:
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(final_md)
    print(f"Markdown file saved successfully: {output_path}")
except Exception as e:
    print("Error saving file:", e)
