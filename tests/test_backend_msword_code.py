"""Code-block detection tests for the DOCX backend."""

from pathlib import Path

from docling_core.types.doc import CodeItem, ContentLayer
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

from docling.backend.msword_backend import MsWordDocumentBackend
from docling.datamodel.base_models import InputFormat
from docling.datamodel.document import DoclingDocument, InputDocument, TextItem

FIXTURE = Path("./tests/data/docx/sources/docx_code_blocks.docx")


def _convert(path) -> DoclingDocument:
    in_doc = InputDocument(
        path_or_stream=path, format=InputFormat.DOCX, backend=MsWordDocumentBackend
    )
    return in_doc._backend.convert()


def _convert_built(document, tmp_path) -> DoclingDocument:
    path = tmp_path / "case.docx"
    document.save(str(path))
    return _convert(path)


def _add_mono(document, text: str, font: str = "Consolas"):
    para = document.add_paragraph()
    para.add_run(text).font.name = font


def _add_code_style(document, name: str = "Source Code", font: str | None = None):
    style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if font is not None:
        style.font.name = font
    return style


def _code_items(doc: DoclingDocument) -> list[CodeItem]:
    return [it for it in doc.texts if isinstance(it, CodeItem)]


def _plain_texts(doc: DoclingDocument) -> set[str]:
    # CodeItem is a TextItem subclass, so it is excluded explicitly rather
    # than relying on isinstance(item, TextItem) alone.
    return {
        item.text
        for item in doc.texts
        if isinstance(item, TextItem) and not isinstance(item, CodeItem)
    }


def test_docx_code_blocks():
    """Style-name and monospaced-font signals, consecutive-paragraph merging,
    and the negative/false-positive guards on the committed fixture.
    """
    doc = _convert(FIXTURE)

    code_items = _code_items(doc)
    assert len(code_items) == 3, (
        "Expected 3 CodeItems: 'Source Code' style, font-only, and the merged "
        "multi-line font-only block"
    )

    assert code_items[0].text == "import sys\nprint(sys.argv)", (
        "Consecutive 'Source Code'-styled paragraphs should merge into one CodeItem"
    )
    assert code_items[1].text == "SELECT * FROM users WHERE active = 1;"
    assert code_items[2].text == (
        "def fib(n):\n    a, b = 0, 1\n    for _ in range(n):\n"
        "        a, b = b, a + b\n    return a"
    ), "Consecutive monospaced paragraphs should merge with indentation preserved"

    negative_texts = {
        "Call the printf function to print formatted output to standard out.",
        "See the original source for details.",
        "Listing 3.2",
        "This memo is set in a typewriter face for a vintage look and feel throughout.",
    }
    assert negative_texts <= _plain_texts(doc), (
        "Negative cases must remain plain TextItems, not CodeItems"
    )


def test_docx_code_block_element_boundaries(tmp_path):
    """Merging must stop at intervening elements, preserve interior blank
    lines, drop trailing blank paragraphs, and treat a 1x1 furniture table as
    a block boundary.
    """
    # An intervening picture breaks the block and preserves document order.
    png = tmp_path / "dot.png"
    Image.new("RGB", (32, 32), (10, 20, 30)).save(str(png))
    d = Document()
    _add_mono(d, "x = 1;")
    d.add_picture(str(png))
    _add_mono(d, "y = 2;")
    doc = _convert_built(d, tmp_path)
    assert [c.text for c in _code_items(doc)] == ["x = 1;", "y = 2;"], (
        "Code blocks must not merge across an intervening picture"
    )
    seq = [
        type(it).__name__
        for it, _ in doc.iterate_items()
        if isinstance(it, CodeItem) or type(it).__name__ == "PictureItem"
    ]
    assert seq == ["CodeItem", "PictureItem", "CodeItem"], (
        "Document order must be code, picture, code"
    )

    # A blank line inside a code block is preserved.
    d = Document()
    _add_code_style(d)
    for line in ("line1 = 1", "", "line3 = 3"):
        d.add_paragraph(line, style="Source Code")
    doc = _convert_built(d, tmp_path)
    assert [c.text for c in _code_items(doc)] == ["line1 = 1\n\nline3 = 3"]

    # Trailing blank code paragraphs do not leave empty lines in the block.
    d = Document()
    _add_code_style(d)
    d.add_paragraph("x = 1;", style="Source Code")
    d.add_paragraph("", style="Source Code")
    d.add_paragraph("", style="Source Code")
    d.add_paragraph("The end (prose).")
    doc = _convert_built(d, tmp_path)
    assert [c.text for c in _code_items(doc)] == ["x = 1;"], (
        "Trailing blank code paragraphs must not append blank lines"
    )

    # A 1x1 furniture table is a boundary: code before/after does not merge
    # with code inside the cell, and a blank code paragraph next to the table
    # must not "spend" the boundary barrier.
    d = Document()
    _add_code_style(d)
    d.add_paragraph("before = 1;", style="Source Code")
    onecell = d.add_table(rows=1, cols=1)
    onecell.rows[0].cells[0].paragraphs[0].style = "Source Code"
    onecell.rows[0].cells[0].paragraphs[0].add_run("inside = 2;")
    d.add_paragraph("", style="Source Code")  # blank code paragraph
    d.add_paragraph("after = 3;", style="Source Code")
    doc = _convert_built(d, tmp_path)
    assert [c.text for c in _code_items(doc)] == [
        "before = 1;",
        "inside = 2;",
        "after = 3;",
    ]


def test_docx_code_style_name_negatives(tmp_path):
    """Style names that merely contain 'code', and ambiguous caption-style
    names, must not be flagged by the style tier.
    """
    d = Document()
    for name in ("Barcode", "Unicode", "Area Code", "QR Code"):
        d.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    d.add_paragraph("Scan the barcode to register the product.", style="Barcode")
    d.add_paragraph("Unicode support was added in version two.", style="Unicode")
    d.add_paragraph("Enter your area code in the box.", style="Area Code")
    d.add_paragraph("Scan the QR code to continue.", style="QR Code")
    doc = _convert_built(d, tmp_path)
    assert not _code_items(doc), "Substring 'code' style names must not be flagged"

    # A style literally named 'Listing' is commonly the CAPTION style in
    # report templates; it must not be treated as an authoritative code style.
    d = Document()
    d.styles.add_style("Listing", WD_STYLE_TYPE.PARAGRAPH)
    d.add_paragraph("Listing 3.1: The main event loop of the server.", style="Listing")
    doc = _convert_built(d, tmp_path)
    assert not _code_items(doc), "'Listing' caption styles are not code styles"


def test_docx_code_font_negatives(tmp_path):
    """The monospace-font fallback must not classify monospaced prose as code:
    neither via everyday punctuation, nor via a monospace document-default
    font, nor for indented or post-code prose lines.
    """
    # A monospaced prose line right after real code stays text.
    d = Document()
    _add_mono(d, "total = a + b;")
    _add_mono(d, "End of the worked example.")
    doc = _convert_built(d, tmp_path)
    assert [c.text for c in _code_items(doc)] == ["total = a + b;"]
    assert "End of the worked example." in _plain_texts(doc), (
        "A non-code monospaced line after code must remain text"
    )

    # Isolated indented monospaced prose (no active code block) stays text.
    d = Document()
    d.add_paragraph("An introductory sentence.")
    _add_mono(d, "    an indented monospaced remark")
    doc = _convert_built(d, tmp_path)
    assert not _code_items(doc), "Indented monospaced prose alone is not code"

    # Everyday parentheses/brackets are not code-indicative: monospaced
    # letterheads, bracket-checkbox form lines, and parenthesized prose all
    # stay text even with explicit run-level monospace fonts.
    d = Document()
    _add_mono(d, "Phone: (555) 123-4567", font="Courier New")
    _add_mono(d, "Fax: (555) 765-4321", font="Courier New")
    _add_mono(d, "[ ] Approved    [ ] Rejected    [ ] Needs review", font="Courier New")
    _add_mono(
        d,
        "Please review the attached memo (see appendix A) before Friday.",
        font="Courier New",
    )
    doc = _convert_built(d, tmp_path)
    assert not _code_items(doc), (
        "Phone numbers, bracket checkboxes and parenthesized prose are not code"
    )

    # A document whose DEFAULT style is monospaced (screenplays, typewriter
    # memos) is not a code document: the default font carries no author
    # intent. An explicit monospace run in the same document still counts.
    d = Document()
    d.styles["Normal"].font.name = "Courier New"
    d.add_paragraph("JOHN (V.O.)")
    d.add_paragraph("(beat) I never expected the results (see Table 3) to hold.")
    d.add_paragraph("Payment terms: net thirty (30) days from receipt of invoice.")
    para = d.add_paragraph()
    para.add_run("total = a + b;").font.name = "Consolas"
    doc = _convert_built(d, tmp_path)
    assert [c.text for c in _code_items(doc)] == ["total = a + b;"], (
        "The document-default font must not count as code evidence; explicit "
        "run-level monospace fonts still do"
    )

    # A custom typewriter-style contract: semicolon-only punctuation,
    # subsection references, plural markers and citation brackets are all
    # prose shapes, not statements or calls.
    d = Document()
    tw = d.styles.add_style("Typewriter", WD_STYLE_TYPE.PARAGRAPH)
    tw.font.name = "Courier New"
    for clause in (
        "AGREEMENT made this first day of January (the Effective Date).",
        "1. The Party of the First Part (hereinafter Seller) agrees;",
        "2. The Party of the Second Part (hereinafter Buyer) agrees;",
        "Refer to Section 12(b) and paragraph 3(c) of the Agreement.",
        "The party(ies) shall return the executed form(s) to the clerk.",
        "He came; he saw; he conquered the entire realm.",
        "The measured effect[12] agrees with earlier work[3].",
        "Includes the following item(s):",
        "The undersigned party(ies):",
        "Enclosed item(s):",
        "Attached form(s):",
        "Responsible party(ies):",
    ):
        d.add_paragraph(clause, style="Typewriter")
    doc = _convert_built(d, tmp_path)
    assert not _code_items(doc), (
        "Typewriter-styled legal and literary prose shapes are not code"
    )

    # A proportional-font tracked insertion (w:ins) is part of the paragraph
    # text, so it must break the all-monospace check.
    d = Document()
    para = d.add_paragraph()
    para.add_run("x = 1;").font.name = "Consolas"
    ins = OxmlElement("w:ins")
    ins.set(qn("w:id"), "1")
    ins.set(qn("w:author"), "Reviewer")
    ins_run = OxmlElement("w:r")
    ins_text = OxmlElement("w:t")
    ins_text.text = "  and a long proportional-font prose insertion about budgets"
    ins_text.set(qn("xml:space"), "preserve")
    ins_run.append(ins_text)
    ins.append(ins_run)
    para._p.append(ins)
    doc = _convert_built(d, tmp_path)
    assert not _code_items(doc), (
        "A proportional-font tracked insertion must break the monospace check"
    )


def test_docx_code_list_interactions(tmp_path):
    """Code and lists must not corrupt each other: lists close before code,
    list items are never absorbed into blocks, and a list that resumes after
    a code block must not fuse the surrounding blocks.
    """
    # A list directly followed by code must close the list; the code and
    # everything after it must not stay nested inside the ListGroup.
    d = Document()
    d.add_paragraph("First bullet", style="List Bullet")
    _add_code_style(d, font="Consolas")
    d.add_paragraph("x = 1;", style="Source Code")
    d.add_paragraph("A normal paragraph after the code block.")
    doc = _convert_built(d, tmp_path)
    codes = _code_items(doc)
    assert len(codes) == 1
    assert type(codes[0].parent.resolve(doc)).__name__ != "ListGroup", (
        "Code after a list must not be nested inside the ListGroup"
    )
    trailing = next(
        it
        for it in doc.texts
        if isinstance(it, TextItem)
        and not isinstance(it, CodeItem)
        and it.text.startswith("A normal paragraph")
    )
    assert type(trailing.parent.resolve(doc)).__name__ != "ListGroup", (
        "Content after a list-then-code sequence must not stay inside the list"
    )

    # A monospaced, code-like LIST item after a code block stays a list item.
    d = Document()
    _add_mono(d, "x = 1;")
    for item_text in ("config.set();", "another();"):
        li = d.add_paragraph(style="List Bullet")
        li.add_run(item_text).font.name = "Consolas"
    doc = _convert_built(d, tmp_path)
    assert [c.text for c in _code_items(doc)] == ["x = 1;"]
    assert {"config.set();", "another();"} <= {it.text for it in doc.texts}, (
        "Monospaced list items must survive as their own items, not be absorbed"
    )

    # A code-STYLED paragraph carrying a stray numPr must close the list and
    # emit code at body level (not nested inside the ListGroup).
    d = Document()
    _add_code_style(d)
    d.add_paragraph("First bullet", style="List Bullet")
    code_para = d.add_paragraph("x = 1;", style="Source Code")
    p_pr = code_para._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_pr.append(ilvl)
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(num_id)
    p_pr.append(num_pr)
    d.add_paragraph("Second bullet", style="List Bullet")
    doc = _convert_built(d, tmp_path)
    codes = _code_items(doc)
    assert len(codes) == 1
    assert type(codes[0].parent.resolve(doc)).__name__ != "ListGroup", (
        "A code paragraph with a stray numPr must not nest inside the ListGroup"
    )

    # A list that RESUMES (same numId) after a code block reuses its cached
    # ListGroup without appending a new body child; the next code block must
    # still start fresh instead of fusing with the pre-list block.
    d = Document()
    _add_code_style(d)
    d.add_paragraph("bullet one", style="List Bullet")
    d.add_paragraph("bullet two", style="List Bullet")
    d.add_paragraph("code_a = 1;", style="Source Code")
    d.add_paragraph("bullet three", style="List Bullet")
    d.add_paragraph("bullet four", style="List Bullet")
    d.add_paragraph("code_b = 2;", style="Source Code")
    doc = _convert_built(d, tmp_path)
    assert [c.text for c in _code_items(doc)] == ["code_a = 1;", "code_b = 2;"], (
        "Code blocks must not fuse across a resumed list"
    )
    assert {"bullet one", "bullet two", "bullet three", "bullet four"} <= {
        it.text for it in doc.texts
    }, "All list items must survive the code interruptions"


def test_docx_code_style_inheritance_and_cells(tmp_path):
    """Style-tier detection must follow base_style inheritance and style-level
    fonts, and must work inside table cells without corrupting tables.
    """
    # A code style inherited via base_style is detected.
    d = Document()
    base = _add_code_style(d)
    child = d.styles.add_style("Project Listing", WD_STYLE_TYPE.PARAGRAPH)
    child.base_style = base
    d.add_paragraph("value = lookup(key);", style="Project Listing")
    doc = _convert_built(d, tmp_path)
    assert len(_code_items(doc)) == 1, "Style inheriting from a code style is code"

    # A custom code style whose monospace font is defined at the STYLE level
    # (so run.font.name is None) is detected via the effective style font.
    d = Document()
    fw = d.styles.add_style("Fixed Width", WD_STYLE_TYPE.PARAGRAPH)
    fw.font.name = "Consolas"
    for line in ("server {", "  listen 80;", "}"):
        d.add_paragraph(line, style="Fixed Width")
    doc = _convert_built(d, tmp_path)
    assert [c.text for c in _code_items(doc)] == ["server {\n  listen 80;\n}"]

    # The style font resolves through the base_style chain: a style without
    # its own font inherits the monospace font of its parent style.
    d = Document()
    mono_base = d.styles.add_style("Mono Base", WD_STYLE_TYPE.PARAGRAPH)
    mono_base.font.name = "Consolas"
    mono_sub = d.styles.add_style("Mono Sub", WD_STYLE_TYPE.PARAGRAPH)
    mono_sub.base_style = mono_base
    d.add_paragraph("y = compute(2);", style="Mono Sub")
    doc = _convert_built(d, tmp_path)
    assert [c.text for c in _code_items(doc)] == ["y = compute(2);"], (
        "A style inheriting a monospace font via base_style is detected"
    )

    # A style with no font anywhere in its chain contributes nothing; runs
    # with explicit monospace fonts still count on their own.
    d = Document()
    d.styles.add_style("Plain Note", WD_STYLE_TYPE.PARAGRAPH)
    para = d.add_paragraph(style="Plain Note")
    para.add_run("x = 1;").font.name = "Consolas"
    doc = _convert_built(d, tmp_path)
    assert [c.text for c in _code_items(doc)] == ["x = 1;"], (
        "Explicit run fonts count even when the style chain has no font"
    )

    # A monospaced style applied to prose (no code punctuation) stays text.
    d = Document()
    memo = d.styles.add_style("Memo Mono", WD_STYLE_TYPE.PARAGRAPH)
    memo.font.name = "Courier New"
    d.add_paragraph(
        "This memo uses a typewriter face for a vintage look and feel.",
        style="Memo Mono",
    )
    doc = _convert_built(d, tmp_path)
    assert not _code_items(doc), "Monospaced prose without code punctuation is text"

    # A code-styled single-paragraph table cell yields a CodeItem.
    d = Document()
    _add_code_style(d, font="Consolas")
    table = d.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Example"
    code_cell = table.rows[0].cells[1]
    code_cell.paragraphs[0].style = "Source Code"
    code_cell.paragraphs[0].add_run("x = compute(y);")
    doc = _convert_built(d, tmp_path)
    assert len(_code_items(doc)) == 1, "A code-styled table cell should emit a CodeItem"

    # An empty code-styled table cell must not crash or corrupt the table.
    d = Document()
    _add_code_style(d)
    table = d.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Label"
    table.rows[0].cells[1].paragraphs[0].style = "Source Code"  # empty
    doc = _convert_built(d, tmp_path)
    assert len(doc.tables) == 1
    assert {c.text for c in doc.tables[0].data.table_cells} == {"Label", ""}


def test_docx_code_furniture_isolation(tmp_path):
    """Header/footer code stays in the furniture layer: it must not merge
    into a body code block nor leak into the body markdown export.
    """
    d = Document()
    _add_code_style(d, font="Consolas")
    d.add_heading("Setup", level=1)
    d.add_paragraph("a = 1;", style="Source Code")
    header = d.sections[0].header
    header.paragraphs[0].style = d.styles["Source Code"]
    header.paragraphs[0].add_run("hdr_line = 99;")
    footer = d.sections[0].footer
    footer.paragraphs[0].style = d.styles["Source Code"]
    footer.paragraphs[0].add_run("ftr_line = 77;")
    doc = _convert_built(d, tmp_path)

    body_codes = [c for c in _code_items(doc) if c.content_layer == ContentLayer.BODY]
    assert [c.text for c in body_codes] == ["a = 1;"], (
        "Header/footer code must not merge into the body code block"
    )
    furniture_codes = [
        c.text for c in _code_items(doc) if c.content_layer == ContentLayer.FURNITURE
    ]
    assert furniture_codes == ["hdr_line = 99;", "ftr_line = 77;"], (
        "Header and footer code are separate furniture-layer CodeItems"
    )
    md = doc.export_to_markdown()
    assert "hdr_line" not in md and "ftr_line" not in md, (
        "Furniture code must stay out of the body markdown export"
    )


def test_docx_code_style_does_not_swallow_checkbox(tmp_path):
    """A checkbox control that happens to carry a code style keeps its
    checkbox label and does not leak the raw glyph into a CodeItem.
    """
    document = Document(str(Path("./tests/data/docx/sources/docx_checkboxes.docx")))
    document.styles.add_style("Source Code", WD_STYLE_TYPE.PARAGRAPH)
    next(
        p for p in document.paragraphs if p.text.strip() == "Design"
    ).style = "Source Code"
    doc = _convert_built(document, tmp_path)

    assert not _code_items(doc), "A code-styled checkbox must not become a CodeItem"
    design = next(it for it in doc.texts if it.text == "Design")
    assert str(getattr(design.label, "value", design.label)).startswith("checkbox"), (
        "The checkbox label and stripped glyph must be preserved"
    )
