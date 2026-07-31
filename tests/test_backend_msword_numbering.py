"""Regression tests for DOCX heading numbering resolution.

Kept in a dedicated module so that ``tests/test_backend_msword.py`` stays under
the repository's max-lines hook.
"""

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docling.backend.msword_backend import MsWordDocumentBackend
from docling.datamodel.base_models import InputFormat
from docling.datamodel.document import InputDocument


def _set_style_numpr(style, num_id=None, ilvl=None):
    """Attach a ``w:numPr`` (numId and/or ilvl) to a paragraph style."""
    num_pr = OxmlElement("w:numPr")
    if ilvl is not None:
        ilvl_elem = OxmlElement("w:ilvl")
        ilvl_elem.set(qn("w:val"), str(ilvl))
        num_pr.append(ilvl_elem)
    if num_id is not None:
        num_id_elem = OxmlElement("w:numId")
        num_id_elem.set(qn("w:val"), str(num_id))
        num_pr.append(num_id_elem)
    style.element.get_or_add_pPr().append(num_pr)


def test_get_numId_and_ilvl_inherits_numid_via_based_on(tmp_path):
    """A heading style carrying only ``ilvl`` inherits ``numId`` from its base.

    Word's stock ``heading 2`` overrides the list level (``ilvl``) but takes the
    ``numId`` from ``heading 1`` through the ``basedOn`` chain. Resolving
    numbering from the paragraph's own style element alone drops that inherited
    numId, so the heading is not recognized as numbered and loses its section
    number, while ``heading 1``/``heading 3`` (which name the numId directly)
    keep theirs. See issue #3916.
    """
    doc = Document()

    # A numbering definition whose levels all render a visible decimal marker.
    numbering = doc.part.numbering_part.element
    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), "77")
    for ilvl in range(3):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(ilvl))
        numfmt = OxmlElement("w:numFmt")
        numfmt.set(qn("w:val"), "decimal")
        lvl.append(numfmt)
        lvltext = OxmlElement("w:lvlText")
        lvltext.set(qn("w:val"), ".".join(f"%{i + 1}" for i in range(ilvl + 1)))
        lvl.append(lvltext)
        abstract_num.append(lvl)
    numbering.append(abstract_num)
    num_elem = OxmlElement("w:num")
    num_elem.set(qn("w:numId"), "42")
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), "77")
    num_elem.append(abstract_ref)
    numbering.append(num_elem)

    base = doc.styles.add_style("Heading A", WD_STYLE_TYPE.PARAGRAPH)
    _set_style_numpr(base, num_id=42, ilvl=0)
    derived = doc.styles.add_style("Heading B", WD_STYLE_TYPE.PARAGRAPH)
    derived.base_style = base
    _set_style_numpr(derived, ilvl=1)  # only ilvl; numId inherited from base

    doc.add_paragraph("Parent section").style = base
    doc.add_paragraph("Child section").style = derived

    docx_path = tmp_path / "based_on_numbering.docx"
    doc.save(str(docx_path))
    backend = InputDocument(
        path_or_stream=docx_path,
        format=InputFormat.DOCX,
        backend=MsWordDocumentBackend,
    )._backend

    paragraphs = {p.text: p for p in backend.docx_obj.paragraphs}
    parent = paragraphs["Parent section"]
    child = paragraphs["Child section"]

    assert backend._get_numId_and_ilvl(parent) == (42, 0)
    # numId inherited from Heading A; before the fix this resolved to (None, 1).
    assert backend._get_numId_and_ilvl(child) == (42, 1)
    assert backend._is_numbered_heading(child)
