from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from docling.document_converter import DocumentConverter


def _add_numbering_definition(doc: Document, abstract_id: str, num_id: str) -> None:
    numbering_part = doc.part.numbering_part
    numbering_xml = numbering_part._element

    abstract_num = OxmlElement("w:abstractNum")
    abstract_num.set(qn("w:abstractNumId"), abstract_id)

    lvl0 = OxmlElement("w:lvl")
    lvl0.set(qn("w:ilvl"), "0")
    start0 = OxmlElement("w:start")
    start0.set(qn("w:val"), "3")
    lvl0.append(start0)
    num_fmt0 = OxmlElement("w:numFmt")
    num_fmt0.set(qn("w:val"), "decimal")
    lvl0.append(num_fmt0)
    lvl_text0 = OxmlElement("w:lvlText")
    lvl_text0.set(qn("w:val"), "%1")
    lvl0.append(lvl_text0)
    abstract_num.append(lvl0)

    lvl1 = OxmlElement("w:lvl")
    lvl1.set(qn("w:ilvl"), "1")
    start1 = OxmlElement("w:start")
    start1.set(qn("w:val"), "1")
    lvl1.append(start1)
    num_fmt1 = OxmlElement("w:numFmt")
    num_fmt1.set(qn("w:val"), "decimal")
    lvl1.append(num_fmt1)
    lvl_text1 = OxmlElement("w:lvlText")
    lvl_text1.set(qn("w:val"), "%1.%2")
    lvl1.append(lvl_text1)
    abstract_num.append(lvl1)

    numbering_xml.insert(0, abstract_num)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), num_id)
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), abstract_id)
    num.append(abstract_num_id)
    numbering_xml.append(num)


def _add_num_pr_to_paragraph(paragraph, num_id: str, ilvl: str) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl_elem = OxmlElement("w:ilvl")
    ilvl_elem.set(qn("w:val"), ilvl)
    num_pr.append(ilvl_elem)
    num_id_elem = OxmlElement("w:numId")
    num_id_elem.set(qn("w:val"), num_id)
    num_pr.append(num_id_elem)
    ppr.append(num_pr)


def _add_num_pr_to_style(style, num_id: str, ilvl: str) -> None:
    style_element = style._element
    ppr = style_element.find("w:pPr", namespaces=style_element.nsmap)
    if ppr is None:
        ppr = OxmlElement("w:pPr")
        style_element.append(ppr)
    num_pr = OxmlElement("w:numPr")
    ilvl_elem = OxmlElement("w:ilvl")
    ilvl_elem.set(qn("w:val"), ilvl)
    num_pr.append(ilvl_elem)
    num_id_elem = OxmlElement("w:numId")
    num_id_elem.set(qn("w:val"), num_id)
    num_pr.append(num_id_elem)
    ppr.append(num_pr)


def _build_docx(path: Path, use_style_inheritance: bool) -> None:
    doc = Document()
    _add_numbering_definition(doc, abstract_id="1", num_id="10")

    if use_style_inheritance:
        base_style = doc.styles["Heading 2"]
        _add_num_pr_to_style(base_style, num_id="10", ilvl="1")
        derived_style = doc.styles.add_style("Heading 2a", WD_STYLE_TYPE.PARAGRAPH)
        derived_style.base_style = base_style
        style = derived_style
        apply_num_pr = False
    else:
        style = doc.styles["Heading 2"]
        apply_num_pr = True

    headings = ["Introduction", "Solution Overview", "Requirements"]
    for heading in headings:
        p = doc.add_paragraph()
        p.style = style
        if apply_num_pr:
            _add_num_pr_to_paragraph(p, num_id="10", ilvl="1")
        p.add_run(heading)

    doc.save(str(path))


def _convert_to_markdown(docx_path: Path) -> str:
    converter = DocumentConverter()
    result = converter.convert(docx_path)
    return result.document.export_to_markdown()


def _assert_numbering(md: str) -> None:
    assert "3.1 Introduction" in md
    assert "3.2 Solution Overview" in md
    assert "3.3 Requirements" in md


def test_docx_numbering_paragraph_numpr(tmp_path: Path) -> None:
    docx_path = tmp_path / "scenario_a.docx"
    _build_docx(docx_path, use_style_inheritance=False)
    md = _convert_to_markdown(docx_path)
    _assert_numbering(md)


def test_docx_numbering_style_inheritance(tmp_path: Path) -> None:
    docx_path = tmp_path / "scenario_b.docx"
    _build_docx(docx_path, use_style_inheritance=True)
    md = _convert_to_markdown(docx_path)
    _assert_numbering(md)
