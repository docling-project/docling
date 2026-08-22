"""Tests for Word table cells wrapped in a content control (``w:sdt``).

Kept separate from ``test_backend_msword.py`` so that file stays under the
repository's per-file line limit.
"""

from docx import Document
from docx.oxml.ns import qn

from docling.datamodel.base_models import InputFormat
from docling.document_converter import DocumentConverter


def _wrap_cell_in_content_control(table, row: int, col: int) -> None:
    """Move a cell under ``w:sdt``/``w:sdtContent``, the way Word does.

    Word produces this shape for date pickers and for cells bound to document
    properties: the ``w:tc`` is no longer a direct child of the ``w:tr``.
    """
    tc = table.cell(row, col)._tc
    tr = tc.getparent()
    position = list(tr).index(tc)
    sdt = tr.makeelement(qn("w:sdt"), {})
    sdt_content = tr.makeelement(qn("w:sdtContent"), {})
    tr.remove(tc)
    sdt_content.append(tc)
    sdt.append(sdt_content)
    tr.insert(position, sdt)


def test_table_cells_inside_a_content_control_keep_their_grid_column(tmp_path):
    """A content-control cell must be parsed, and must not shift the row left.

    Only direct ``w:tc`` children of a ``w:tr`` used to be visited, so a cell
    Word had wrapped in a content control was dropped -- and because the grid
    column advances once per emitted cell, every later cell in the row moved
    into the vacated column and lined up under the wrong header.
    """

    converter = DocumentConverter(allowed_formats=[InputFormat.DOCX])
    header = ["Date", "Version", "Author", "Note"]
    row = ["1.2.2015", "1", "Acme s.r.o.", "Created"]

    def build(wrapped_columns: tuple[int, ...]) -> list[str]:
        doc = Document()
        table = doc.add_table(rows=2, cols=len(header))
        for col, value in enumerate(header):
            table.cell(0, col).text = value
        for col, value in enumerate(row):
            table.cell(1, col).text = value
        for col in wrapped_columns:
            _wrap_cell_in_content_control(table, 1, col)

        path = tmp_path / f"content_control_{len(wrapped_columns)}.docx"
        doc.save(str(path))

        data = converter.convert(path).document.tables[0].data
        grid = [""] * data.num_cols
        for cell in data.table_cells or []:
            if cell.start_row_offset_idx == 1:
                grid[cell.start_col_offset_idx] = cell.text or ""
        return grid

    assert build(()) == row
    # A single wrapped cell: previously dropped, shifting the rest one column left.
    assert build((0,)) == row
    # Two wrapped cells in the same row, one of them not the first.
    assert build((0, 2)) == row
