from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

# Create document with multi-level outline numbering
doc = Document()
numbering_part = doc.part.numbering_part
numbering_xml = numbering_part._element

# Create abstract numbering definition
abstractNum = OxmlElement('w:abstractNum')
abstractNum.set(qn('w:abstractNumId'), '1')

# Level 0 (Chapter - "3")
lvl0 = OxmlElement('w:lvl')
lvl0.set(qn('w:ilvl'), '0')
start0 = OxmlElement('w:start')
start0.set(qn('w:val'), '3')
lvl0.append(start0)
numFmt0 = OxmlElement('w:numFmt')
numFmt0.set(qn('w:val'), 'decimal')
lvl0.append(numFmt0)
lvlText0 = OxmlElement('w:lvlText')
lvlText0.set(qn('w:val'), '%1')
lvl0.append(lvlText0)
abstractNum.append(lvl0)

# Level 1 (Section - "3.1", "3.2")
lvl1 = OxmlElement('w:lvl')
lvl1.set(qn('w:ilvl'), '1')
start1 = OxmlElement('w:start')
start1.set(qn('w:val'), '1')
lvl1.append(start1)
numFmt1 = OxmlElement('w:numFmt')
numFmt1.set(qn('w:val'), 'decimal')
lvl1.append(numFmt1)
lvlText1 = OxmlElement('w:lvlText')
lvlText1.set(qn('w:val'), '%1.%2')
lvl1.append(lvlText1)
abstractNum.append(lvl1)

numbering_xml.insert(0, abstractNum)

num = OxmlElement('w:num')
num.set(qn('w:numId'), '1')
abstractNumId = OxmlElement('w:abstractNumId')
abstractNumId.set(qn('w:val'), '1')
num.append(abstractNumId)
numbering_xml.append(num)

def add_numbered_heading(doc, text, level, num_id):
    p = doc.add_paragraph()
    p.style = f'Heading {level}'
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level - 1))
    numPr.append(ilvl)
    numId_elem = OxmlElement('w:numId')
    numId_elem.set(qn('w:val'), str(num_id))
    numPr.append(numId_elem)
    pPr.append(numPr)
    run = p.add_run(text)
    return p

add_numbered_heading(doc, "Introduction", 2, 1)
doc.add_paragraph("Content under 3.1")
add_numbered_heading(doc, "Solution Overview", 2, 1)
doc.add_paragraph("Content under 3.2")
add_numbered_heading(doc, "Requirements", 2, 1)
doc.add_paragraph("Content under 3.3")

doc.save("/tmp/test_outline_numbering.docx")

# Now parse with Docling
from docling.document_converter import DocumentConverter

converter = DocumentConverter()
result = converter.convert("/tmp/test_outline_numbering.docx")

for item in result.document.iterate_items():
    element = item[0] if isinstance(item, tuple) else item
    label = element.label.value if hasattr(element.label, 'value') else str(element.label)
    text = element.text if hasattr(element, 'text') else ""
    print(f"{label}: {text}")