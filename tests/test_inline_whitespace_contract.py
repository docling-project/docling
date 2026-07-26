"""Backend-side checks for the InlineGroup whitespace contract.

Inline runs carry their own significant whitespace and the docling-core serializers
concatenate them faithfully, with no separator of their own. The per-format boundary cases
live next to each backend's own tests; this module holds the cases that need to build a
source document on the fly.
"""

import re
from pathlib import Path

import pytest
from docx import Document

from docling.backend.html_backend import HTMLDocumentBackend
from docling.backend.msword_backend import MsWordDocumentBackend
from docling.datamodel.base_models import InputFormat
from docling.datamodel.document import InputDocument, TextItem


def test_docx_paragraph_runs_keep_their_own_whitespace(tmp_path):
    """A DOCX run's boundary whitespace stays in the run, not invented by the serializer.

    In DOCX the space at a formatting boundary very often sits *inside* the formatted run.
    Stripping each run forced the serializers to re-insert a separator, which produced
    ``**bold** .`` and ``H 2 O``. Runs are now emitted verbatim, with only the paragraph
    edges trimmed, and the serializer hoists edge whitespace out of the emphasis markers.
    """
    docx_path = tmp_path / "runs.docx"
    document = Document()
    paragraph = document.add_paragraph()
    paragraph.add_run("   Normal ")  # leading paragraph padding is trimmed
    bold = paragraph.add_run("bold ")  # the boundary space lives inside the bold run
    bold.bold = True
    paragraph.add_run("tail.   ")  # trailing paragraph padding is trimmed

    no_space = document.add_paragraph()
    no_space.add_run("H")
    sub = no_space.add_run("2")
    sub.font.subscript = True
    no_space.add_run("O")

    document.save(docx_path)

    in_doc = InputDocument(
        path_or_stream=docx_path,
        format=InputFormat.DOCX,
        backend=MsWordDocumentBackend,
    )
    doc = MsWordDocumentBackend(in_doc=in_doc, path_or_stream=docx_path).convert()

    runs = [item.text for item, _ in doc.iterate_items() if isinstance(item, TextItem)]
    assert runs == ["Normal ", "bold ", "tail.", "H", "2", "O"]

    md = doc.export_to_markdown()
    assert "Normal **bold** tail." in md
    assert "H2O" in md


def test_docx_inline_equations_keep_the_boundary_space():
    """Text after an inline equation keeps its leading space.

    ``_add_inline_equations_to_parent`` builds its own InlineGroup, outside the corrected
    paragraph path, and used to strip the fragment following the last equation. With a
    faithful join that merged the formula into the next word
    (``$A= \\pi r^{2}$is the area formula``).
    """
    source = Path("tests/data/docx/sources/equations.docx")
    in_doc = InputDocument(
        path_or_stream=source,
        format=InputFormat.DOCX,
        backend=MsWordDocumentBackend,
    )
    doc = MsWordDocumentBackend(in_doc=in_doc, path_or_stream=source).convert()
    md = doc.export_to_markdown()

    # list items (the helper's own InlineGroup) ...
    assert (
        "- First item with inline equation: $A= \\pi r^{2}$ is the area formula." in md
    )
    assert (
        "- Second item with equations: $E=mc^{2}$ and $F=ma$ are physics formulas."
        in md
    )
    assert "- The formula $a^{2}+b^{2}=c^{2}$ is the Pythagorean theorem." in md
    # ... and a normal paragraph through the same helper, where the boundary is a period.
    assert (
        "This is a word document and this is an inline equation: $A= \\pi r^{2}$." in md
    )

    # No equation is glued to the following word anywhere in the document.
    assert not re.search(r"\}\$[A-Za-z]", md), md


def _convert_html(tmp_path, body: str):
    source = tmp_path / "t.html"
    source.write_text(f"<html><body>{body}</body></html>", encoding="utf-8")
    in_doc = InputDocument(
        path_or_stream=source,
        format=InputFormat.HTML,
        backend=HTMLDocumentBackend,
    )
    return HTMLDocumentBackend(in_doc=in_doc, path_or_stream=source).convert()


@pytest.mark.parametrize(
    ("body", "expected_md", "expected_txt"),
    [
        # Adjacent code spans: the boundary belongs between the two `` ` `` delimiters, not
        # inside either of them. It used to land inside the second span (``a`` b``).
        ("<p><code>a</code> <code>b</code></p>", "`a` `b`", "a b"),
        ("<p><code>a</code> <code>b</code> <code>c</code></p>", "`a` `b` `c`", "a b c"),
        # No source whitespace means no boundary to preserve.
        ("<p>x<code>a</code><code>b</code>y</p>", "x`ab`y", "xaby"),
        # One side plain: the space goes on the plain run.
        ("<p><code>a</code> <b>b</b></p>", "`a` **b**", "a b"),
        ("<p><b>a</b> <code>b</code></p>", "**a** `b`", "a b"),
        ("<p>see <code>a</code> then</p>", "see `a` then", "see a then"),
        # The original contract cases stay correct.
        (
            "<p>Water is H<sub>2</sub>O here</p>",
            "Water is H2O here",
            "Water is H2O here",
        ),
        (
            "<p>x<sup>2</sup> + y<sup>3</sup> uses sup</p>",
            "x2 + y3 uses sup",
            "x2 + y3 uses sup",
        ),
    ],
)
def test_html_boundary_stays_outside_code_delimiters(
    tmp_path, body, expected_md, expected_txt
):
    doc = _convert_html(tmp_path, body)
    assert doc.export_to_markdown() == expected_md
    assert doc.export_to_text() == expected_txt


def test_html_pre_keeps_interior_whitespace(tmp_path):
    """`<pre>` is preformatted: only the block edges are trimmed, never each fragment.

    Stripping every fragment merged mixed content into one word once the serializer stopped
    inserting a separator. The one-code-item-per-fragment shape of `<pre>` with inline markup
    is a separate pre-existing issue and is not addressed here.
    """
    doc = _convert_html(tmp_path, '<pre>See <a href="x">link</a> first</pre>')
    assert doc.export_to_text() == "See link first"

    plain = _convert_html(tmp_path, "<pre>\ndef f():\n    return 1\n</pre>")
    assert "def f():\n    return 1" in plain.export_to_markdown()
