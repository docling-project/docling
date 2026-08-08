"""Tests for content controls (``w:sdt``) that also contain a picture.

Kept separate from ``test_backend_msword.py`` so that file stays under the
repository's per-file line limit.
"""

from docx import Document
from docx.shared import Inches
from lxml import etree
from PIL import Image

from docling.datamodel.base_models import InputFormat
from docling.document_converter import DocumentConverter

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _move_into_content_control(doc, paragraphs) -> None:
    """Wrap the given body paragraphs in a ``w:sdt``, as Word cover pages are."""
    body = doc.element.body
    sdt = etree.SubElement(body, f"{{{W_NS}}}sdt")
    etree.SubElement(sdt, f"{{{W_NS}}}sdtPr")
    sdt_content = etree.SubElement(sdt, f"{{{W_NS}}}sdtContent")
    for paragraph in paragraphs:
        body.remove(paragraph._p)
        sdt_content.append(paragraph._p)
    body.remove(sdt)
    body.insert(0, sdt)


def test_content_control_text_survives_a_picture_in_the_same_control(tmp_path):
    """A picture inside a content control must not swallow the control's text.

    The image branches of the element walk are keyed on descendant XPaths, so a
    ``w:sdt`` holding a picture anywhere inside used to match there and the
    content-control branch was never reached -- the picture was emitted and
    every paragraph in the control was dropped. Word's built-in cover pages are
    exactly this shape, so the document title went missing.
    """

    converter = DocumentConverter(allowed_formats=[InputFormat.DOCX])
    logo_path = tmp_path / "logo.png"
    Image.new("RGB", (120, 120), (200, 30, 30)).save(str(logo_path))

    def build(with_picture: bool) -> str:
        doc = Document()
        paragraphs = []
        if with_picture:
            picture_paragraph = doc.add_paragraph()
            picture_paragraph.add_run().add_picture(str(logo_path), width=Inches(1.5))
            paragraphs.append(picture_paragraph)
        paragraphs.append(doc.add_paragraph("COVER TITLE INSIDE SDT"))
        _move_into_content_control(doc, paragraphs)
        doc.add_paragraph("BODY TEXT OUTSIDE SDT")

        path = tmp_path / f"content_control_picture_{with_picture}.docx"
        doc.save(str(path))
        return converter.convert(path).document.export_to_markdown()

    without_picture = build(with_picture=False)
    assert "COVER TITLE INSIDE SDT" in without_picture
    assert "BODY TEXT OUTSIDE SDT" in without_picture

    with_picture = build(with_picture=True)
    assert "COVER TITLE INSIDE SDT" in with_picture
    assert "BODY TEXT OUTSIDE SDT" in with_picture
    # The picture is still emitted, and exactly once.
    assert with_picture.count("<!-- image -->") == 1
